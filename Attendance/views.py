from django.shortcuts import render, redirect
from django.utils.html import escape
from . import forms
from django.views.generic import TemplateView
from django.contrib.auth import logout, authenticate, login
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
import json
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.authtoken.models import Token

from rest_framework import generics, status
from .models import Teacher, Student, Lecture, Div, Subject, AppUser
from .models import SubjectTeacher, StudentLecture, StudentDivision
from .serializers import (TeacherSerializer, StudentSerializer, LectureSerializer, DivSerializer, SubjectSerializer,
                          ShortTeacherSerializer)
from rest_framework.authentication import TokenAuthentication
import datetime
import csv


class HomePage(TemplateView):
    template_name = 'Attendance/index.html'


class ThanksPage(TemplateView):
    template_name = 'Attendance/logout_success.html'


def login_user_teacher(request):
    # logout(request)
    # username = password = ''
    if request.POST:
        username = request.POST['username']
        password = request.POST['password']

        user = authenticate(username=username, password=password)
        if user is not None:
            if user.is_active:
                login(request, user)
                return redirect('Attendance:dash')
    return render(request, 'Attendance/login.html', context={'form': forms.TeacherLoginForm()})


@login_required
def dash(request):
    return render(request, 'Attendance/login_success.html')


def signup(request):
    if request.method == 'POST':
        form = forms.UserCreateForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data.get('username')
            raw_password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=raw_password)
            login(request, user)
            return redirect('Attendance:dash')
    else:
        form = forms.UserCreateForm()
    return render(request, 'Attendance/signup.html', {'form': form})


# WEB EndPoint Views


class TeachersSubjectDataView(generics.GenericAPIView):
    authentication_classes = (TokenAuthentication,)
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):

        teacherId = kwargs['teacherId']

        try:
            teacher = Teacher.objects.get(teacherID=teacherId)
        except Exception as e:
            response_data = {'error_message': "Invalid TeacherId" + str(e)}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        divisions = Div.objects.filter(classteacher=teacher, calendar_year=datetime.date.today().year)
        class_subjects = []
        for div in divisions:
            class_subjects_ordereddict = list(set([st.subject for st in SubjectTeacher.objects.filter(div=div)]))
            for subject in class_subjects_ordereddict:
                subject_json = SubjectSerializer(subject).data
                subject_json["div"] = str(div)
                class_subjects.append(subject_json)

        taught_subjects = []

        subjectteacher = SubjectTeacher.objects.filter(teacher=teacher)
        for st in subjectteacher:
            subject_json = SubjectSerializer(st.subject).data
            subject_json["div"] = str(st.div)
            taught_subjects.append(subject_json)

        division = None
        for div in divisions:
            if div.get_class_type() == "Class":
                division = div

        response_data = {
            'taught_subjects': taught_subjects,
            'class_subjects': class_subjects,
            'division_they_are_class_teacher_of': DivSerializer(division).data,
        }

        return JsonResponse(response_data, status=status.HTTP_200_OK)


class LoginTeacherView(generics.GenericAPIView):
    permission_classes = (AllowAny,)

    def post(self, request, *args, **kwargs):

        try:
            form_data = json.loads(request.body.decode())
            teacherId = escape(form_data['teacherId'])
            password = escape(form_data['password'])
        except Exception:
            teacherId = escape(request.POST.get('teacherId'))
            password = escape(request.POST.get('password'))

        teacher = Teacher.objects.get(teacherID=teacherId)
        user = authenticate(username=teacher.user.username, password=password)

        if user is not None:
            token, _ = Token.objects.get_or_create(user=user)
            login(request, user)

            response_data = {
                'token': token.key,
            }

            return JsonResponse(response_data, status=status.HTTP_200_OK)
        else:
            response_data = {'error_message': "Cannot log you in"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)


class LogoutTeacherView(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request):
        # simply delete the token to force a login
        request.user.auth_token.delete()
        logout(request)
        response_data = {'success_message': 'Successfully logged you out'}
        return JsonResponse(response_data, status=status.HTTP_200_OK)


class SignUpTeacherView(generics.GenericAPIView):
    permission_classes = (AllowAny,)

    def post(self, request, *args, **kwargs):

        try:
            form_data = json.loads(request.body.decode())
            teacherId = escape(form_data['teacherId'])
            password = escape(form_data['password'])
            f_name = escape(form_data['fname'])
            l_name = escape(form_data['lname'])
            specialization = escape(form_data['spec'])
        except Exception:
            teacherId = escape(request.POST.get('teacherId'))
            password = escape(request.POST.get('password'))
            f_name = escape(request.POST.get('fname'))
            l_name = escape(request.POST.get('lname'))
            specialization = escape(request.POST.get('spec'))

        try:
            user = AppUser.objects.create(username=teacherId, password=password)
            user.first_name = f_name
            user.last_name = l_name
            user.set_password(password)
            user.is_teacher = True
            user.save()
            teacher = Teacher.objects.create(user=user, teacherID=teacherId, specialization=specialization)
            teacher.save()

            login(request, user)

            token, _ = Token.objects.get_or_create(user=user)

            response_data = {
                'token': token.key,
            }
            return JsonResponse(response_data, status=status.HTTP_200_OK)
        except Exception as e:
            response_data = {'error_message': "Cannot sign you up due to " + str(e)}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)


class GenericLoginView(generics.GenericAPIView):
    permission_classes = (AllowAny,)

    def post(self, request, *args, **kwargs):

        try:
            form_data = json.loads(request.body.decode())
            user_id = escape(form_data['id'])
            password = escape(form_data['password'])
        except Exception:
            user_id = escape(request.POST.get('id'))
            password = escape(request.POST.get('password'))

        user = authenticate(username=user_id, password=password)

        if user is not None:
            token, _ = Token.objects.get_or_create(user=user)
            login(request, user)

            response_data = {
                'isStudent': user.is_student,
                'isTeacher': user.is_teacher,
                'token': token.key,
            }

            if user.is_student:
                response_data['user'] = StudentSerializer(Student.objects.get(user=user)).data
            else:
                response_data['user'] = ShortTeacherSerializer(Teacher.objects.get(user=user)).data

            response_data['user']['name'] = user.getname()

            return JsonResponse(response_data, status=status.HTTP_200_OK)
        else:
            response_data = {'error_message': "Cannot log you in"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)


class GetAttendanceOfDay(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):

        subject_name = kwargs['subject']
        div = kwargs['div']
        try:
            date = kwargs['date']
            d, m, y = date.split('-')
            date = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date = datetime.date.today()

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)
        if date.month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        teacher = Teacher.objects.get(user=request.user)

        if div.classteacher is teacher:
            lecs = Lecture.objects.filter(date=date, div=div, subject=subject, attendanceTaken=True)
        else:
            lecs = Lecture.objects.filter(date=date, teacher=teacher, div=div, subject=subject, attendanceTaken=True)

        attendance_list = {}

        for lec in lecs:
            lecTime = lec.getTimeString()
            attendance_list[lecTime] = []
            student_lecs = StudentLecture.objects.filter(lecture=lec)
            present_students = [sl.student for sl in student_lecs]

            student_divs = StudentDivision.objects.filter(division=div)
            div_students = [sd.student for sd in student_divs]

            absent_students = list(set(div_students) - set(present_students))

            for student in present_students:
                student_json = StudentSerializer(student).data
                student_json["attendance"] = 1
                attendance_list[lecTime].append(student_json)

            for student in absent_students:
                student_json = StudentSerializer(student).data
                student_json["attendance"] = 0
                attendance_list[lecTime].append(student_json)

            attendance_list[lecTime].sort(key=lambda x: x["sapID"])

        final_attendance_list = []

        for lecTime in attendance_list:
            attendance_object = {}
            attendance_object['time'] = lecTime
            attendance_object['attendance_list'] = attendance_list[lecTime]
            final_attendance_list.append(attendance_object)

        response_data = {
            'attendance': final_attendance_list,
        }

        return JsonResponse(response_data, status=status.HTTP_200_OK)


class GetAttendanceOfRange(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):

        subject_name = kwargs['subject']
        div = kwargs['div']

        try:
            date_from = kwargs['date_from']
            d, m, y = date_from.split('-')
            date_from = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_from = datetime.date.today()

        try:
            date_to = kwargs['date_to']
            d, m, y = date_to.split('-')
            date_to = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_to = datetime.date.today()

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)

        if date_from.month < 6 and date_to.month < 6:
            semester = year * 2
        elif date_from.month >= 6 and date_to.month >= 6:
            semester = year * 2 - 1
        else:
            response_data = {'error_message': "Dates are not from the same semester."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        teacher = Teacher.objects.get(user=request.user)

        if div.classteacher is teacher:
            lecs = Lecture.objects.filter(date__lte=date_to, date__gte=date_from, div=div, subject=subject,
                                          attendanceTaken=True)
        else:
            lecs = Lecture.objects.filter(date__lte=date_to, date__gte=date_from, teacher=teacher,
                                          div=div, subject=subject, attendanceTaken=True)

        attendance_list = {}

        for lec in lecs:
            lecDateTime = lec.getDateTimeString()
            attendance_list[lecDateTime] = []
            student_lecs = StudentLecture.objects.filter(lecture=lec)
            present_students = [sl.student for sl in student_lecs]

            student_divs = StudentDivision.objects.filter(division=div)
            div_students = [sd.student for sd in student_divs]

            absent_students = list(set(div_students) - set(present_students))

            for student in present_students:
                student_json = StudentSerializer(student).data
                student_json["attendance"] = 1
                attendance_list[lecDateTime].append(student_json)

            for student in absent_students:
                student_json = StudentSerializer(student).data
                student_json["attendance"] = 0
                attendance_list[lecDateTime].append(student_json)

            attendance_list[lecDateTime].sort(key=lambda x: x["sapID"])

        final_attendance_list = []

        for lecDateTime in attendance_list:
            attendance_object = {}
            attendance_object['time'] = lecDateTime
            attendance_object['attendance_list'] = attendance_list[lecDateTime]
            final_attendance_list.append(attendance_object)

        response_data = {
            'attendance': final_attendance_list,
        }

        return JsonResponse(response_data, status=status.HTTP_200_OK)


class GetAttendanceOfStudent(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def multiple_lectures(self, lecture):

        if lecture.div.get_class_type() == 'Practical':
            return 1

        start = datetime.datetime.combine(lecture.date, lecture.startTime)
        end = datetime.datetime.combine(lecture.date, lecture.endTime)
        difference = end - start
        td = difference.total_seconds() / 60
        if td > 90 and td <= 150:
            return 2
        elif td > 150 and td < 210:
            return 3
        return 1

    def get(self, request, *args, **kwargs):

        subject_name = kwargs['subject']
        student_sap = int(kwargs['sapID'])

        try:
            subject = Subject.objects.get(name=subject_name)
            student = Student.objects.get(sapID=student_sap)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Student.DoesNotExist:
            response_data = {'error_message': "Student with SAP " + str(student_sap) + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        divs = list(student.div.all())

        class_teacher_list = [div.classteacher for div in divs]

        teacher = Teacher.objects.get(user=request.user)

        if teacher in class_teacher_list:
            lecs = Lecture.objects.filter(div__in=divs, subject=subject, attendanceTaken=True)
        else:
            lecs = Lecture.objects.filter(teacher=teacher, div__in=divs, subject=subject, attendanceTaken=True)

        if lecs:
            lecs = list(lecs)
            lecs.sort(key=lambda x: x.date)

            attendance_list = []
            attendance_count = 0
            attendance_total = 0

            for lecture in lecs:
                lecture_json = LectureSerializer(lecture).data
                lecture_json["date"] = "-".join(lecture_json["date"].split('-')[::-1])
                try:
                    StudentLecture.objects.get(student=student, lecture=lecture)
                    lecture_json["attendance"] = 1
                    attendance_count += self.multiple_lectures(lecture)
                except StudentLecture.DoesNotExist:
                    lecture_json["attendance"] = 0

                attendance_total += self.multiple_lectures(lecture)
                attendance_list.append(lecture_json)

            attendance_percentage = attendance_count * 100 / attendance_total

            response_data = {
                'attendance': attendance_list,
                'attendance_count': attendance_count,
                'attendance_total': attendance_total,
                'attendance_percentage': attendance_percentage,
            }

        else:
            response_data = {
                'attendance': [],
            }

        return JsonResponse(response_data, status=status.HTTP_200_OK)


class EditAttendanceOfDay(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):

        subject_name = kwargs['subject']
        div = kwargs['div']

        try:
            form_data = json.loads(request.body.decode())
            attendance_list = form_data['attendance_list']
        except Exception:
            attendance_list = request.POST.get('attendance_list')

        try:
            date = kwargs['date']
            d, m, y = date.split('-')
            date = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date = datetime.date.today()

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)
        if date.month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        teacher = Teacher.objects.get(user=request.user)

        if div.classteacher is teacher:
            lecs = Lecture.objects.filter(date=date, div=div, subject=subject)
        else:
            lecs = Lecture.objects.filter(date=date, teacher=teacher, div=div, subject=subject)

        for attendance_object in attendance_list:
            current_lecture = None
            lecTime = attendance_object['time']
            for lec in lecs:
                if lec.getTimeString() == lecTime:
                    current_lecture = lec

            for student_entry in attendance_object['attendance_list']:
                student = Student.objects.get(sapID=student_entry['sapID'])
                if int(student_entry['attendance']) == 1:
                    StudentLecture.objects.get_or_create(student=student, lecture=current_lecture)
                else:
                    try:
                        sl = StudentLecture.objects.get(student=student, lecture=current_lecture)
                        sl.delete()
                    except StudentLecture.DoesNotExist:
                        pass

        response_data = {'success_message': 'Successfully saved attendance data'}
        return JsonResponse(response_data, status=status.HTTP_200_OK)


class DownloadCsv(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def multiple_lectures(self, lecture):

        if lecture.div.get_class_type() == 'Practical':
            return 1

        start = datetime.datetime.combine(lecture.date, lecture.startTime)
        end = datetime.datetime.combine(lecture.date, lecture.endTime)
        difference = end - start
        td = difference.total_seconds() / 60
        if td > 90 and td <= 150:
            return 2
        elif td > 150 and td < 210:
            return 3
        return 1

    def get(self, request, *args, **kwargs):
        subject_name = kwargs['subject']
        div = kwargs['div']

        try:
            date_from = kwargs['date_from']
            d, m, y = date_from.split('-')
            date_from = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_from = datetime.date.today()

        try:
            date_to = kwargs['date_to']
            d, m, y = date_to.split('-')
            date_to = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_to = datetime.date.today()

        teacher = Teacher.objects.get(user=request.user)

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)

        if date_from.month < 6 and date_to.month < 6:
            semester = year * 2
        elif date_from.month > 6 and date_to.month > 6:
            semester = year * 2 - 1
        else:
            response_data = {'error_message': "Dates are not from the same semester."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            SubjectTeacher.objects.get(div=div, subject=subject, teacher=teacher)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except SubjectTeacher.DoesNotExist:
            response_data = {'error_message': "You do not have access to " + subject_name + " for " + div + " data."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        lecs = Lecture.objects.filter(date__lte=date_to, date__gte=date_from, div=div, subject=subject,
                                      attendanceTaken=True)

        total = 0
        for lec in lecs:
            count = self.multiple_lectures(lec)
            total += count

        student_list = Student.objects.filter(div=div).order_by('sapID')
        student_lectures = StudentLecture.objects.filter(lecture__in=lecs)
        attendance_list = []
        for student in student_list:
            relevant_student_lectures = student_lectures.filter(student=student)

            student_attended = 0
            for lec in relevant_student_lectures:
                count = self.multiple_lectures(lec.lecture)
                student_attended += count

            student_json = StudentSerializer(student).data
            student_json["attendance_count"] = student_attended
            if lecs:
                student_json["attendance_percentage"] = student_attended * 100 / total
            else:
                student_json["attendance_percentage"] = 100
            attendance_list.append(student_json)

        attendance_list.sort(key=lambda x: x["sapID"])

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'blob; filename="AttendanceData.csv"'

        csvwriter = csv.writer(response)

        csvwriter.writerow(["SAP ID", "Name", "Attendance Count (" + str(total) + ")", "Attendance Percentage"])
        for student in attendance_list:
            csvwriter.writerow([student["sapID"], student["name"], student["attendance_count"],
                                student["attendance_percentage"]])

        return response


class DownloadSAPSheet(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        subject_name = kwargs['subject']
        div = kwargs['div']

        try:
            date_from = kwargs['date_from']
            d, m, y = date_from.split('-')
            date_from = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_from = datetime.date.today()

        try:
            date_to = kwargs['date_to']
            d, m, y = date_to.split('-')
            date_to = datetime.datetime(int(y), int(m), int(d)).date()
        except KeyError:
            date_to = datetime.date.today()

        teacher = Teacher.objects.get(user=request.user)

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)

        if date_from.month < 6 and date_to.month < 6:
            semester = year * 2
        elif date_from.month >= 6 and date_to.month >= 6:
            semester = year * 2 - 1
        else:
            response_data = {'error_message': "Dates are not from the same semester."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            SubjectTeacher.objects.get(div=div, subject=subject, teacher=teacher)

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + division + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except SubjectTeacher.DoesNotExist:
            response_data = {
                'error_message': "You do not have access to " + subject_name + " for " + str(div) + " data."
            }
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        is_practical = div.get_class_type() == "Practical"

        if is_practical:
            divs = Div.objects.filter(
                division__contains=div.division[0],
                division__regex=r'[A,B][1-4]',
                semester=semester,
                calendar_year=datetime.date.today().year
            )

            lecs = Lecture.objects.filter(date__gte=date_from, date__lte=date_to, div__in=divs, subject=subject,
                                          attendanceTaken=True).order_by('date')
            student_list = Student.objects.filter(div__in=divs).order_by('sapID')
            student_divs = StudentDivision.objects.filter(division__in=divs, student__in=student_list)

        else:
            lecs = Lecture.objects.filter(date__gte=date_from, date__lte=date_to, div=div, subject=subject,
                                          attendanceTaken=True).order_by('date')
            student_list = Student.objects.filter(div=div).order_by('sapID')
            student_divs = StudentDivision.objects.filter(division=div, student__in=student_list)

        student_lectures = StudentLecture.objects.filter(lecture__in=lecs)

        attendance_sheet = []

        for student in student_list:
            student_row = [student.sapID, str(student).upper()]
            for lec in lecs:
                if student_lectures.filter(lecture=lec, student=student).exists():
                    student_row.append('P')
                elif student_divs.filter(student=student, division=lec.div).exists():
                    student_row.append('Ab')
            attendance_sheet.append(student_row)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'blob; filename="WeeklyAttendance.docx"'

        document = Document()

        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 2
        style = document.styles['Table Grid']
        font = style.font
        font.name = 'Cambria'
        font.size = Pt(11)

        try:
            document.add_picture('/home/wizdem/Attendance-System-Web/SAP/header.png', width=Inches(6))
        except FileNotFoundError:
            document.add_picture('SAP/header.png', width=Inches(6))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        digits = div.calendar_year % 100
        if div.semester % 2 == 0:
            academic_year = str(div.calendar_year - 1) + "-" + str(digits)
        else:
            academic_year = str(div.calendar_year) + "-" + str(digits + 1)

        p = document.add_paragraph('Academic Year: ' + academic_year)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('Report of Attendance Record of Lectures')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=5, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.cell(0, 0).text = "Week No.:"
        table.cell(0, 1).text = "Date: "
        table.cell(0, 2).text = "From: " + date_from.strftime("%d-%m-%Y")
        table.cell(0, 3).text = "To:" + date_to.strftime("%d-%m-%Y")
        table.cell(1, 0).merge(table.cell(1, 1)).text = "Teacher: " + (str(teacher) if not is_practical else "")
        table.cell(1, 2).merge(table.cell(1, 3)).text = "Subject: " + subject.name
        table.cell(2, 0).text = "Class: " + yearname
        table.cell(2, 1).text = "Course/Branch: Computer"
        table.cell(2, 2).text = "Semester: " + str(semester)
        table.cell(2, 3).text = "Division: " + division[0]
        table.cell(3, 0).merge(table.cell(3, 1)).text = "No. of Lectures Scheduled (S):"
        table.cell(3, 2).merge(table.cell(3, 3)).text = "No. of Lectures Conducted (C):"
        table.cell(4, 0).merge(table.cell(4, 3)).text = "Remark (In case S ≠ C):"

        p = document.add_paragraph('')

        if is_practical:
            section = document.add_section(WD_SECTION.CONTINUOUS)

            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), str(len(divs)))

            for i in range(1, 5):
                prac_div = divs.get(division=division[0] + str(i))
                prac_list = []

                for row in attendance_sheet:
                    stu_list = student_list.filter(sapID=row[0])
                    if student_divs.filter(student__in=stu_list, division=prac_div).exists():
                        prac_list.append(row)

                lec_list = lecs.filter(div=prac_div)

                table = document.add_table(rows=4, cols=1 + (len(lec_list) if (len(lec_list) > 1) else 1))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                table.cell(0, 0).merge(table.cell(0, 1)).text = "Batch: " + str(prac_div)
                table.cell(1, 0).merge(table.cell(3, 0)).text = "SAP ID"

                if len(lec_list) <= 1:
                    table.cell(1, 1).text = "Date"
                else:
                    table.cell(1, 1).merge(table.cell(1, len(lec_list))).text = "Date"

                col_no = 1
                for lec in lecs:
                    table.cell(2, col_no).text = str(lec.date.strftime("%d/%m"))
                    table.cell(3, col_no).text = str(lec.getShortTimeString())
                    col_no += 1

                for entry in prac_list:
                    row = table.add_row()
                    row.cells[0].text = str(entry[0])
                    for i, val in enumerate(entry[2:]):
                        row.cells[i + 1].text = str(val)

                if i < 4:
                    document.add_section(WD_SECTION.NEW_COLUMN)

        else:
            if(len(lecs) > 6):
                cols = 3 + len(lecs)
            else:
                cols = 9

            table = document.add_table(rows=3 + len(attendance_sheet), cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            for col in table.columns:
                col.width = Inches(0.6)
                for cell in col.cells:
                    cell.width = Inches(0.6)
            col = table.columns[0]
            col.width = Inches(0.4)
            for cell in col.cells:
                cell.width = Inches(0.4)
            col = table.columns[1]
            col.width = Inches(1.1)
            for cell in col.cells:
                cell.width = Inches(1.1)
            col = table.columns[2]
            col.width = Inches(2)
            for cell in col.cells:
                cell.width = Inches(2)

            p = table.cell(0, 0).merge(table.cell(2, 0)).paragraphs[0]
            p.text = "Sr.No."
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = table.cell(0, 1).merge(table.cell(2, 1)).paragraphs[0]
            p.text = "SAP ID"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = table.cell(0, 2).merge(table.cell(2, 2)).paragraphs[0]
            p.text = "Name"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if(len(lecs) > 6):
                p = table.cell(0, 3).merge(table.cell(0, 2 + len(lecs))).paragraphs[0]
            else:
                p = table.cell(0, 3).merge(table.cell(0, 8)).paragraphs[0]

            p.text = "Date & Time"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            col_no = 3
            for lec in lecs:
                table.cell(1, col_no).text = str(lec.date.strftime("%d/%m"))
                table.cell(2, col_no).text = str(lec.getShortTimeString())
                col_no += 1

            row_no = 3
            for row in attendance_sheet:
                table.cell(row_no, 0).text = str(row_no - 2)
                col_no = 1
                for val in row:
                    table.cell(row_no, col_no).text = str(val)
                    col_no += 1
                row_no += 1

        document.save(response)

        return response

# Android EndPoint Views


class GetLectureListOfTheDay(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        try:
            date = kwargs['date']
            d, m, y = date.split('-')
            date = datetime.datetime(int(y), int(m), int(d)).date()
        except Exception:
            response_data = {'error_message': "Date sent is incorrectly formatted. Excepted dd-mm-yy format."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        try:
            teacher = Teacher.objects.get(user=request.user)
        except Exception:
            response_data = {'error_message': "Logged in user is not a teacher."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        week_day = date.weekday() + 2
        if week_day == 8:
            week_day = 1

        all_past_lectures = Lecture.objects.filter(teacher=teacher, date__week_day=week_day)

        subjectdivs = []

        lectures = []

        predicted_lectures = []

        if datetime.date.today().month < 6:
            rem = 0
        else:
            rem = 1

        for lec in all_past_lectures:
            div_is_correct = (lec.div.semester % 2 == rem) and (lec.div.calendar_year == datetime.date.today().year)
            if (lec.subject, lec.div) not in subjectdivs and div_is_correct:
                subjectdivs.append((lec.subject, lec.div))

        for subjectdiv in subjectdivs:
            pastlecs = all_past_lectures.filter(subject=subjectdiv[0], div=subjectdiv[1])
            max = None
            counts = {}
            for lec in pastlecs:
                if lec.startTime in counts:
                    counts[lec.startTime] += 1
                else:
                    counts[lec.startTime] = 1

                if not max or counts[lec.startTime] > counts[max.startTime]:
                    max = lec
            if max:
                lectures.append(max)

            for lec in pastlecs:
                if lec.date == date and lec.startTime in counts and max.startTime != lec.startTime:
                    lectures.append(lec)

        for ttlecture in sorted(lectures, key=lambda x: x.startTime):
            try:
                lecture = Lecture.objects.get(
                    roomNumber=ttlecture.roomNumber,
                    startTime=ttlecture.startTime,
                    endTime=ttlecture.endTime,
                    date=date,
                    teacher=ttlecture.teacher,
                    div=ttlecture.div,
                    subject=ttlecture.subject
                )
                if ttlecture.date == date:
                    predicted = 0
                else:
                    predicted = 1

            except Lecture.DoesNotExist:
                lecture = Lecture(
                    roomNumber=ttlecture.roomNumber,
                    startTime=ttlecture.startTime,
                    endTime=ttlecture.endTime,
                    date=date,
                    teacher=ttlecture.teacher,
                    div=ttlecture.div,
                    subject=ttlecture.subject
                )
                predicted = 1

            lecture_json = LectureSerializer(lecture).data
            lecture_json['type'] = lecture.div.get_class_type()
            lecture_json['attendanceTaken'] = 1 if lecture.attendanceTaken else 0
            lecture_json['predicted'] = predicted
            predicted_lectures.append(lecture_json)

        return JsonResponse({
            'date': date,
            'lectures': predicted_lectures
        }, status=status.HTTP_200_OK)


class GetStudentListOfLecture(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        subject_name = kwargs['subject']
        div = kwargs['div']
        lecture_date = kwargs['date']
        startTime = kwargs['startTime']

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)
        if datetime.date.today().month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            h, m, s = startTime.split(':')
            startTime = datetime.time(int(h), int(m), int(s))
            d, m, y = lecture_date.split('-')
            lec_date = datetime.datetime(int(y), int(m), int(d)).date()

        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Exception:
            response_data = {'error_message': "Wrong Date and/or Time format. Expecting dd-mm-yy and hh:mm:ss"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        students = Student.objects.filter(div=div).order_by('sapID')
        students_json = StudentSerializer(students, many=True).data

        teacher = Teacher.objects.get(user=request.user)
        try:
            lecture = Lecture.objects.get(subject=subject, div=div, date=lec_date, teacher=teacher, startTime=startTime)

            for student in students_json:
                student_object = students.get(sapID=student['sapID'])
                student['name'] = student_object.user.getname()
                try:
                    StudentLecture.objects.get(lecture=lecture, student=student_object)
                    student['Attendance'] = 1
                except StudentLecture.DoesNotExist:
                    student['Attendance'] = 0
                student['sapID'] = str(student['sapID'])

        except Lecture.DoesNotExist:
            for student in students_json:
                student_object = students.get(sapID=student['sapID'])
                student['Attendance'] = 0
                student['sapID'] = str(student['sapID'])
                student['name'] = student_object.user.getname()

        return JsonResponse({'students': students_json}, status=status.HTTP_200_OK)


class SaveAttendance(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        try:
            form_data = json.loads(request.body.decode())
            subject_name = escape(form_data['subject'])
            div = escape(form_data['div'])
            roomNumber = escape(form_data['room'])
            startTime = escape(form_data['startTime'])
            endTime = escape(form_data['endTime'])
            lecture_date = escape(form_data['date'])
            students = form_data['students']
        except KeyError:
            response_data = {'error_message': "Expecting subject, div, room, startTime, endTime, date and students."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            subject_name = escape(request.POST.get('subject'))
            div = escape(request.POST.get('div'))
            roomNumber = escape(request.POST.get('room'))
            startTime = escape(request.POST.get('startTime'))
            endTime = escape(request.POST.get('endTime'))
            lecture_date = escape(request.POST.get('date'))
            students = request.POST.get('students')

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)

        teacher = Teacher.objects.get(user=request.user)

        if datetime.date.today().month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            SubjectTeacher.objects.get(div=div, subject=subject, teacher=teacher)
            h, m, s = startTime.split(':')
            startTime = datetime.time(int(h), int(m), int(s))
            h, m, s = endTime.split(':')
            endTime = datetime.time(int(h), int(m), int(s))
            d, m, y = lecture_date.split('-')
            lecture_date = datetime.datetime(int(y), int(m), int(d)).date()
        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except SubjectTeacher.DoesNotExist:
            response_data = {'error_message': "Division " + str(div) + " does not have Subject " + subject_name}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Exception:
            response_data = {'error_message': "Wrong Date and/or Time format. Expecting dd-mm-yy and hh:mm:ss"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        lecture, _ = Lecture.objects.get_or_create(
            roomNumber=roomNumber,
            startTime=startTime,
            endTime=endTime,
            date=lecture_date,
            teacher=Teacher.objects.get(user=request.user),
            div=div,
            subject=subject
        )

        sapIDs = [int(student['sapID']) for student in students]
        student_objects = Student.objects.filter(sapID__in=sapIDs).order_by('sapID')

        for student in students:
            student_object = student_objects.get(sapID=int(student['sapID']))
            if student['Attendance'] == 1:
                StudentLecture.objects.get_or_create(student=student_object, lecture=lecture)
            else:
                try:
                    StudentLecture.objects.get(student=student_object, lecture=lecture).delete()
                except StudentLecture.DoesNotExist:
                    pass

        lecture.attendanceTaken = True
        lecture.save()

        return JsonResponse({
            'subject': lecture.subject.name,
            'div': str(lecture.div),
            'room': lecture.roomNumber,
            'startTime': lecture.startTime.strftime('%H:%M:%S'),
            'endTime': lecture.endTime.strftime('%H:%M:%S'),
            'date': lecture.date.strftime('%d-%m-%Y')
        }, status=status.HTTP_200_OK)


class GetStudentsAttendance(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def multiple_lectures(self, lecture):

        if lecture.div.get_class_type() == 'Practical':
            return 1

        start = datetime.datetime.combine(lecture.date, lecture.startTime)
        end = datetime.datetime.combine(lecture.date, lecture.endTime)
        difference = end - start
        td = difference.total_seconds() / 60
        if td > 90 and td <= 150:
            return 2
        elif td > 150 and td < 210:
            return 3
        return 1

    def get(self, request, *args, **kwargs):
        user = request.user
        student = Student.objects.get(user=user)
        divisions = Div.objects.filter(calendar_year=datetime.date.today().year, student=student)

        attendance = {}

        if datetime.date.today().month < 6:
            rem = 0
        else:
            rem = 1

        for division in divisions:
            if division.semester % 2 == rem:
                subjects = [st.subject for st in SubjectTeacher.objects.filter(div=division)]
                for subject in subjects:
                    lectures = Lecture.objects.filter(div=division, subject=subject, attendanceTaken=True)
                    type = division.get_class_type()
                    attendance[subject.name + type] = {
                        'type': type,
                        'subject': subject.name,
                        'total': 0,
                        'attended': 0,
                    }
                    for lec in lectures:
                        count = self.multiple_lectures(lec)
                        attendance[subject.name + type]['total'] += count
                        try:
                            StudentLecture.objects.get(student=student, lecture=lec)
                            attendance[subject.name + type]['attended'] += count
                        except StudentLecture.DoesNotExist:
                            pass

        attendance_list = []
        for sub in attendance:
            attendance_list.append(attendance[sub])

        return JsonResponse({'attendance': attendance_list}, status=status.HTTP_200_OK)


class GetStudentAttendanceHistory(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        subject_name = kwargs['subject']
        type = kwargs['type']

        user = request.user
        student = Student.objects.get(user=user)
        divisions = Div.objects.filter(calendar_year=datetime.date.today().year, student=student)

        try:
            subject = Subject.objects.get(name=subject_name)
        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        if type not in ['Class', 'Practical', 'Elective']:
            response_data = {'error_message': "Type " + type + " does not match 'Class', 'Practical' or 'Elective'."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        lectures = []
        for division in divisions:
            div_has_subject = SubjectTeacher.objects.filter(div=division, subject=subject).exists()
            if division.get_class_type() == type and div_has_subject:
                lectures.extend(list(Lecture.objects.filter(div=division, subject=subject, attendanceTaken=True)))

        lecs_json = []
        lectures = sorted(lectures, key=lambda lec: lec.date, reverse=True)

        for lec in lectures:
            if StudentLecture.objects.filter(lecture=lec, student=student).exists():
                present = 1
            else:
                present = 0
            lecs_json.append({
                'date': lec.date.strftime("%d-%m-%Y"),
                'time': lec.getTimeString(),
                'present': present
            })
        return JsonResponse(lecs_json, status=status.HTTP_200_OK, safe=False)


class GetSubjectsAndDivisions(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def get(self, request, *args, **kwargs):
        teacher = Teacher.objects.get(user=request.user)

        subjectteachers = SubjectTeacher.objects.filter(teacher=teacher)

        return_objects = []

        for st in subjectteachers:
            if st.div.calendar_year == datetime.date.today().year:
                obj = {
                    'div': str(st.div),
                    'subject': st.subject.name
                }
                return_objects.append(obj)

        return JsonResponse(return_objects, status=status.HTTP_200_OK, safe=False)


class SaveLectureAndGetStudentsList(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        try:
            form_data = json.loads(request.body.decode())
            subject_name = escape(form_data['subject'])
            div = escape(form_data['div'])
            roomNumber = escape(form_data['room'])
            startTime = escape(form_data['startTime'])
            endTime = escape(form_data['endTime'])
            lecture_date = escape(form_data['date'])
        except KeyError:
            response_data = {'error_message': "Expecting subject, div, room, startTime, endTime and date."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            subject_name = escape(request.POST.get('subject'))
            div = escape(request.POST.get('div'))
            roomNumber = escape(request.POST.get('room'))
            startTime = escape(request.POST.get('startTime'))
            endTime = escape(request.POST.get('endTime'))
            lecture_date = escape(request.POST.get('date'))

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)

        teacher = Teacher.objects.get(user=request.user)

        if datetime.date.today().month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            SubjectTeacher.objects.get(div=div, subject=subject, teacher=teacher)
            h, m, s = startTime.split(':')
            startTime = datetime.time(int(h), int(m), int(s))
            h, m, s = endTime.split(':')
            endTime = datetime.time(int(h), int(m), int(s))
            d, m, y = lecture_date.split('-')
            lecture_date = datetime.datetime(int(y), int(m), int(d)).date()
        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except SubjectTeacher.DoesNotExist:
            response_data = {'error_message': "Division " + str(div) + " does not have Subject " + subject_name}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Exception:
            response_data = {'error_message': "Wrong Date and/or Time format. Expecting dd-mm-yy and hh:mm:ss"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        lecture, _ = Lecture.objects.get_or_create(
            roomNumber=roomNumber,
            startTime=startTime,
            endTime=endTime,
            date=lecture_date,
            teacher=teacher,
            div=div,
            subject=subject
        )

        students = Student.objects.filter(div=div).order_by('sapID')
        students_json = StudentSerializer(students, many=True).data

        for student in students_json:
            student_object = students.get(sapID=student['sapID'])
            student['name'] = student_object.user.getname()
            try:
                StudentLecture.objects.get(lecture=lecture, student=student_object)
                student['Attendance'] = 1
            except StudentLecture.DoesNotExist:
                student['Attendance'] = 0
            student['sapID'] = str(student['sapID'])

        return JsonResponse({
            'students': students_json
        }, status=status.HTTP_200_OK)


class DeleteLecture(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        try:
            form_data = json.loads(request.body.decode())
            subject_name = escape(form_data['subject'])
            div = escape(form_data['div'])
            roomNumber = escape(form_data['room'])
            startTime = escape(form_data['startTime'])
            endTime = escape(form_data['endTime'])
            if 'date' in form_data:
                lecture_date = escape(form_data['date'])
            else:
                lecture_date = datetime.date.today()
        except KeyError:
            response_data = {'error_message': "Expecting subject, div, room, startTime and endTime. Date optional."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            subject_name = escape(request.POST.get('subject'))
            div = escape(request.POST.get('div'))
            roomNumber = escape(request.POST.get('room'))
            startTime = escape(request.POST.get('startTime'))
            endTime = escape(request.POST.get('endTime'))
            lecture_date = escape(request.POST.get('number', datetime.date.today()))

        teacher = Teacher.objects.get(user=request.user)

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)
        if datetime.date.today().month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            subject = Subject.objects.get(name=subject_name)
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            SubjectTeacher.objects.get(div=div, subject=subject, teacher=teacher)
            h, m, s = startTime.split(':')
            startTime = datetime.time(int(h), int(m), int(s))
            h, m, s = endTime.split(':')
            endTime = datetime.time(int(h), int(m), int(s))
            if isinstance(lecture_date, str):
                d, m, y = lecture_date.split('-')
                lecture_date = datetime.datetime(int(y), int(m), int(d)).date()
        except Subject.DoesNotExist:
            response_data = {'error_message': "Subject " + subject_name + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except SubjectTeacher.DoesNotExist:
            response_data = {'error_message': "Division " + str(div) + " does not have Subject " + subject_name}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Exception:
            response_data = {'error_message': "Wrong Date and/or Time format. Expecting dd-mm-yy and hh:mm:ss"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        try:
            lecture = Lecture.objects.get(
                roomNumber=roomNumber,
                startTime=startTime,
                endTime=endTime,
                date=lecture_date,
                teacher=teacher,
                div=div,
                subject=subject
            )
            lecture.delete()
            deleted = 1
        except Lecture.DoesNotExist:
            deleted = 0
            pass

        return JsonResponse({'success': deleted}, status=status.HTTP_200_OK)


class ChangePassword(generics.GenericAPIView):
    permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        try:
            form_data = json.loads(request.body.decode())
            oldpass = form_data['old_password']
            newpass = form_data['new_password']

        except KeyError:
            response_data = {'error_message': "Expecting old_password and new_password."}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            oldpass = request.POST.get('old_password')
            newpass = request.POST.get('new_password')

        user = request.user
        user = authenticate(username=user.username, password=oldpass)
        if user is not None:
            user.set_password(newpass)
            user.save()
            return JsonResponse({'success': 1}, status=status.HTTP_200_OK)

        return JsonResponse({'success': 0}, status=status.HTTP_200_OK)


class GetPreviousLectureAttendance(generics.GenericAPIView):
    permission_classes = (IsAuthenticated, )

    def get(self, request, *args, **kwargs):
        div = kwargs['div']
        lecture_date = kwargs['date']
        startTime = kwargs['startTime']

        yearname, division = div.split("_")
        year = Div.yearnameToYear(yearname)
        if datetime.date.today().month < 6:
            semester = year * 2
        else:
            semester = year * 2 - 1
        try:
            div = Div.objects.get(division=division, semester=semester, calendar_year=datetime.date.today().year)
            h, m, s = startTime.split(':')
            startTime = datetime.time(int(h), int(m), int(s))
            d, m, y = lecture_date.split('-')
            lec_date = datetime.datetime(int(y), int(m), int(d)).date()

        except Div.DoesNotExist:
            response_data = {'error_message': "Division " + div + " Does Not Exist"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        except Exception:
            response_data = {'error_message': "Wrong Date and/or Time format. Expecting dd-mm-yy and hh:mm:ss"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)

        students = Student.objects.filter(div=div).order_by('sapID')
        students_json = StudentSerializer(students, many=True).data

        try:
            days_lectures = Lecture.objects.filter(div=div, date=lec_date, startTime__lt=startTime)
            lecture = days_lectures.filter(attendanceTaken=True).latest('startTime')
        except Lecture.DoesNotExist:
            response_data = {'error_message': "No previous lectures on this day"}
            return JsonResponse(response_data, status=status.HTTP_400_BAD_REQUEST)
        for student in students_json:
            student_object = students.get(sapID=student['sapID'])
            student['name'] = student_object.user.getname()
            try:
                StudentLecture.objects.get(lecture=lecture, student=student_object)
                student['Attendance'] = 1
            except StudentLecture.DoesNotExist:
                student['Attendance'] = 0
            student['sapID'] = str(student['sapID'])

        return JsonResponse({
            'students': students_json,
            'subject': lecture.subject.name,
            'timing': lecture.getTimeString()
        }, status=status.HTTP_200_OK)
